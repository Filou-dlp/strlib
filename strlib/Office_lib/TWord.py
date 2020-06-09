import docx # for microsoft office

# Openoffice
from odf.opendocument import OpenDocumentText as odt 
from odf.style import (Style, TextProperties, ParagraphProperties, ListLevelProperties, TabStop, TabStops)
from odf.text import (H, P, List, ListItem, ListStyle, ListLevelStyleNumber, ListLevelStyleBullet)
from odf import teletype

class TWord:
    """
        Class to do word document for microsoft office
    """

    def __init__(self,name,office):
        """
            Constructor
            :@parma name: Name of the document
            :@type name: String
        """
        self.__name = name
        self.__office = office
    
    def create(self):
        """
            Function to create the document
        """
        if self.__office == "WORD":
            self.__document = Document()  
        elif self.__office == "ODT":
            self.__document = OpenDocumentText()
    
    def save(self):
        """
            Function to save the document
        """
        if self.__office == "WORD":
            self.__document.save(self.__name+".docx")
        
        elif self.__office == "ODT":
            self.__document.save(self.__name+u".odt")
    
    def close(self):
        """
            Function to close and save the document
        """
        self.save()
        del self.__document
        
    def title(self,text,lvl=0):
        """
            Function to add a title to the document
            :@param text: Text of the title
            :@param lvl: level of the header
            :@type text: string
            :@type lvl: integer
            :@default lvl: 0 (First header)
        """
        
        if self.__office == "WORD":
            self.__document.add_heading(text, level=lvl)
            
        elif self.__office == "ODT":
            h1style_ = Style(name="CenterHeading "+str(lvl+1))
            mymainheading_element_ = H(outlinelevel=lvl, stylename=h1style_)
            teletype.addTextToElement(mymainheading_element_, text)

            self.__document.text.addElement(mymainheading_element_)


    def paragraphe(self,text):
        """
            Function to add a paragraph to the document
            :@param text: Text of the title
            :@type text: string
        """
        if self.__office == "WORD":
            self.__document.add_paragraph(text)
            
        elif self.__office == "ODT": 
            paragraph_element_ = P(stylename=justifystyle)
            teletype.addTextToElement(paragraph_element_, text)

            self.__document.text.addElement(paragraph_element_, text)
        
    def page_break(self):
        """
            Function to save the document
        """
        if self.__office == "WORD":
            self.__document.add_page_break()
        else:
            raise NotImplementedError
    
        
        
if __name__ == '__main__':
    pass
